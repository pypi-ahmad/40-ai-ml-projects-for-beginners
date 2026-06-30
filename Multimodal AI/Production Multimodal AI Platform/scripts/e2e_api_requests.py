"""Execute real HTTP requests against running FastAPI server."""

from __future__ import annotations

import json
import time
from pathlib import Path
from typing import Any

import httpx
from PIL import Image, ImageDraw


def ensure_assets() -> dict[str, str]:
    base = Path("data/uploads")
    base.mkdir(parents=True, exist_ok=True)

    image_a = base / "api_e2e_a.png"
    image_b = base / "api_e2e_b.png"
    docx_path = base / "api_e2e_notes.docx"

    img = Image.new("RGB", (640, 360), color=(245, 248, 252))
    draw = ImageDraw.Draw(img)
    draw.rectangle((40, 40, 260, 170), outline=(30, 80, 120), width=3)
    draw.text((60, 70), "Ops Dashboard", fill=(30, 80, 120))
    draw.text((60, 110), "Revenue +15%", fill=(20, 20, 20))
    img.save(image_a)

    img2 = Image.new("RGB", (640, 360), color=(238, 244, 250))
    draw2 = ImageDraw.Draw(img2)
    draw2.rectangle((44, 46, 264, 176), outline=(20, 70, 110), width=3)
    draw2.text((64, 76), "Ops Dashboard", fill=(20, 70, 110))
    draw2.text((64, 116), "Revenue +17%", fill=(20, 20, 20))
    img2.save(image_b)

    from docx import Document

    doc = Document()
    doc.add_heading("API E2E Notes", level=1)
    doc.add_paragraph("Revenue increased and latency stayed stable.")
    doc.add_paragraph("Error counts reduced after deployment.")
    doc.save(docx_path)

    return {
        "image_a": str(image_a),
        "image_b": str(image_b),
        "docx": str(docx_path),
    }


def wait_ready(base_url: str) -> None:
    last = ""
    for _ in range(120):
        try:
            res = httpx.get(f"{base_url}/health", timeout=2)
            if res.status_code == 200:
                return
        except Exception as exc:  # noqa: BLE001
            last = str(exc)
        time.sleep(0.5)
    raise RuntimeError(f"API not ready: {last}")


def post_ok(base_url: str, path: str, payload: dict[str, Any]) -> dict[str, Any]:
    res = httpx.post(f"{base_url}{path}", json=payload, timeout=60)
    if res.status_code != 200:
        raise RuntimeError(f"{path} status={res.status_code} body={res.text}")
    body = res.json()
    if body.get("status") != "ok":
        raise RuntimeError(f"{path} non-ok payload={body}")
    return body


def main() -> None:
    base_url = "http://127.0.0.1:8000"
    wait_ready(base_url)
    assets = ensure_assets()

    trace = {"trace": {"source": "api_e2e"}}
    responses: dict[str, Any] = {}

    responses["health"] = httpx.get(f"{base_url}/health", timeout=30).json()
    responses["caption"] = post_ok(
        base_url,
        "/caption",
        {**trace, "input": {"image_path": assets["image_a"]}, "options": {"style": "short"}},
    )
    responses["ocr"] = post_ok(
        base_url,
        "/ocr",
        {**trace, "input": {"document_path": assets["docx"]}},
    )
    responses["documents"] = post_ok(
        base_url,
        "/documents",
        {**trace, "input": {"document_path": assets["docx"]}},
    )
    responses["embeddings"] = post_ok(
        base_url,
        "/embeddings",
        {**trace, "input": {"text": "revenue increased"}},
    )
    responses["compare"] = post_ok(
        base_url,
        "/compare",
        {**trace, "input": {"image_paths": [assets["image_a"], assets["image_b"]]}},
    )
    responses["vqa"] = post_ok(
        base_url,
        "/vqa",
        {
            **trace,
            "input": {
                "document_path": assets["docx"],
                "question": "What are key insights from this document?",
            },
        },
    )
    responses["search"] = post_ok(
        base_url,
        "/search",
        {
            **trace,
            "input": {"query": "latency stable"},
            "options": {"modality": "document", "top_k": 5},
        },
    )
    responses["retrieve"] = post_ok(
        base_url,
        "/retrieve",
        {
            **trace,
            "input": {"query": "error reduced"},
            "options": {"modality": "document", "top_k": 5},
        },
    )
    responses["analyze"] = post_ok(
        base_url,
        "/analyze",
        {**trace, "input": {"image_path": assets["image_a"], "question": "Explain this image"}},
    )
    responses["analytics"] = post_ok(base_url, "/analytics", trace)

    report_dir = Path("outputs/reports")
    report_dir.mkdir(parents=True, exist_ok=True)
    out_json = report_dir / "api_e2e_results.json"
    out_md = report_dir / "api_e2e_results.md"

    out_json.write_text(
        json.dumps({"assets": assets, "responses": responses}, indent=2, default=str),
        encoding="utf-8",
    )

    lines = ["# API E2E Results", ""]
    for key, value in responses.items():
        lines.append(
            f"- `{key}`: status={value.get('status')} "
            f"latency_ms={value.get('latency_ms')}"
        )
    out_md.write_text("\n".join(lines), encoding="utf-8")

    print(f"Wrote {out_json}")
    print(f"Wrote {out_md}")


if __name__ == "__main__":
    main()
